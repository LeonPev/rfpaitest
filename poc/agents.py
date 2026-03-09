"""
IDF Procurement AI — Multi-agent pipeline.
Real product: processes actual uploaded documents via Gemini.
Domain-aware, IDF-focused, quality-reviewed output.
"""

import os, json, time, re, unicodedata, io
from pathlib import Path
from datetime import datetime

from google import genai
from google.genai.types import FileState, UploadFileConfig
from dotenv import load_dotenv
import docx  # python-docx

load_dotenv()
_client = genai.Client(api_key=os.getenv('GEMINI_API_KEY'))

BASE_DIR = Path(__file__).parent.parent   # rfpaitest/
MODEL       = 'gemini-2.5-flash'
FLASH_MODEL = 'gemini-2.5-flash'

MIME_MAP = {
    '.doc':  'application/msword',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    '.pdf':  'application/pdf',
    '.txt':  'text/plain',
    '.md':   'text/plain',
}

# Domain context mapping for IDF procurement categories
DOMAIN_CONTEXTS = {
    "דלק": {
        "domain": "דלק ושמנים",
        "regulations": "תקנות משרד האנרגיה, תקני ISO 4259, ASTM D975, EN 590",
        "keywords": ["דלק", "סולר", "בנזין", "שמנים", "סיכה", "תחנת דלק", "מכלית"],
        "expertise": "תדלוק כלי רכב ומערכות צבאיות, ניהול מכליות, תקני בטיחות אחסון דלקים, עמידה בתקנות איכות הסביבה למניעת דליפות",
        "vendors_focus": "חברות דלק ישראליות (פז, סונול, דור-אלון, דלק), ספקי שמנים מיוחדים לצבא",
        "avoid": "אל תמליץ על חברות הייטק, חברות תוכנה, או ספקים שלא עוסקים בתחום הדלקים. התמקד אך ורק בספקי דלק, שמנים, חומרי סיכה ותחזוקה של מערכות דלק."
    },
    "חומרי ניקוי": {
        "domain": "חומרי ניקוי והיגיינה",
        "regulations": "תקנת REACH EC 1907/2006, BPR EU 528/2012, ISO 9001, ISO 14001, SDS",
        "keywords": ["ניקוי", "חיטוי", "היגיינה", "סבון", "אקונומיקה"],
        "expertise": "חומרי ניקוי מוסדיים, תקני בטיחות כימית, גיליונות SDS, דרישות REACH לחומרים כימיים",
        "vendors_focus": "יצרני חומרי ניקוי ישראליים (סנו, שטראוס, כתר), יבואני חומרי ניקוי מקצועיים",
        "avoid": "אל תמליץ על ספקי טכנולוגיה או שירותים שאינם קשורים לתחום חומרי הניקוי. התמקד בספקים שמייצרים או מייבאים חומרי ניקוי."
    },
    "מזון": {
        "domain": "מזון ומוצרי מזון",
        "regulations": "ISO 22000, HACCP, תקנות כשרות (רבנות ראשית), תקנות משרד הבריאות, תקן ישראלי ת\"י",
        "keywords": ["מזון", "כריכים", "פרי", "ירק", "כשרות", "אספקת מזון", "קייטרינג"],
        "expertise": "שרשרת קירור, כשרות צבאית (מהדרין/רגילה), תנאי אחסון מזון, תקנות יבוא מזון, בדיקות מעבדה",
        "vendors_focus": "ספקי מזון מוסדיים, חברות קייטרינג צבאי, יצרני מזון ישראליים בעלי תו תקן",
        "avoid": "אל תמליץ על ספקים שאין להם תעודת כשרות, אל תציע פתרונות טכנולוגיים שלא קשורים לשרשרת המזון."
    },
    "ציוד רפואי": {
        "domain": "ציוד רפואי וכירורגי",
        "regulations": "ISO 13485, CE marking, FDA 510(k), תקנות משרד הבריאות, AMAR",
        "keywords": ["רפואי", "כירורגי", "מתכלה", "אביזרי רפואה", "ציוד הנשמה"],
        "expertise": "ציוד רפואי לשימוש צבאי, רפואת חירום שדה, סטריליזציה, שרשרת אספקה רפואית, תחזוקה מונעת",
        "vendors_focus": "ספקי ציוד רפואי מאושרים (B. Braun, Medtronic, GE Healthcare), יבואנים מורשים",
        "avoid": "אל תמליץ על ספקי ציוד שאינו רפואי. התמקד בציוד בעל אישורי CE/FDA ורישום AMAR."
    },
    "ריהוט": {
        "domain": "ריהוט ופלסטיק",
        "regulations": "EN 16139, EN 1022, תקן ישראלי לריהוט, ISO 7173",
        "keywords": ["ריהוט", "כסאות", "שולחנות", "פלסטיק", "ארגונומי"],
        "expertise": "ריהוט מוסדי/צבאי עמיד, עמידות למזג אוויר, ארגונומיה, עמידה בעומסים",
        "vendors_focus": "יצרני ריהוט ישראליים (כתר, סטארפלסט), ספקי ריהוט מוסדי",
        "avoid": "אל תמליץ על ריהוט ביתי/יוקרתי. התמקד בריהוט מוסדי עמיד ומתאים לתנאי שטח."
    },
    "כלי מטבח": {
        "domain": "כלי מטבח והגשה",
        "regulations": "EU Food Contact Regulation EC 1935/2004, תקנות משרד הבריאות, ISO 8442",
        "keywords": ["מטבח", "כלים", "הגשה", "סכו\"ם", "סירי בישול"],
        "expertise": "כלי מטבח מוסדיים, עמידות לשימוש כבד, מגע עם מזון, כשרות",
        "vendors_focus": "ספקי כלי מטבח מוסדיים (Gastro International), יצרני כלי בישול ישראליים",
        "avoid": "אל תמליץ על כלי מטבח ביתיים. התמקד בציוד מוסדי מקצועי."
    },
}

IDF_SYSTEM_CONTEXT = """הקשר מערכתי חשוב:
- אתה עובד עבור מערכת הביטחון של מדינת ישראל (משהב"ט/צה"ל)
- כל ההמלצות חייבות להתחשב בדרישות ביטחוניות, סיווגים, ובטיחות מידע
- יש להתחשב ב: חוק חובת מכרזים, תקנות משרד הביטחון (DOPP), הנחיות חשב"ד
- עדיפות לספקים ישראליים שיכולים לספק לבסיסי צה"ל
- יש להתחשב בתנאי מצב חירום ומלחמה (הגדלת הספקה, גמישות לוגיסטית)
- מערכת המרק"ט (SRM) היא המערכת הממוחשבת לניהול הזמנות
- בכל SOW חדש יש לכלול: מצב חירום, ביטחון מידע, עיצומים, הבטחת איכות
- דיוק חשוב! עדיף פחות ומדויק מאשר הרבה וגנרי. הימנע מהמלצות כלליות שלא מתאימות לתחום."""


def _detect_domain(file_paths: list, user_context: str) -> dict:
    """Detect the procurement domain from file names and user context."""
    combined = " ".join([Path(p).name for p in file_paths]) + " " + (user_context or "")
    combined_lower = combined.lower()

    best_match = None
    best_score = 0

    for domain_key, ctx in DOMAIN_CONTEXTS.items():
        score = sum(1 for kw in ctx["keywords"] if kw in combined_lower or kw in combined)
        if score > best_score:
            best_score = score
            best_match = domain_key

    if best_match and best_score > 0:
        return DOMAIN_CONTEXTS[best_match]
    return {
        "domain": "כללי",
        "regulations": "חוק חובת מכרזים, תקנות משהב\"ט",
        "keywords": [],
        "expertise": "רכש ממשלתי כללי",
        "vendors_focus": "ספקים מאושרים ע\"י משרד הביטחון",
        "avoid": "הימנע מהמלצות גנריות שלא קשורות לתחום הספציפי."
    }


# ─── Document discovery ───────────────────────────────────────────────────────

def list_sow_docs():
    """Scan data/SOW/ recursively and return structured file list."""
    sow_dir = BASE_DIR / "data" / "SOW"
    docs = []
    if not sow_dir.exists():
        return docs
    for p in sorted(sow_dir.rglob("*")):
        if not p.is_file():
            continue
        if p.suffix.lower() not in MIME_MAP:
            continue
        if p.name.startswith('~$'):
            continue
        if 'fixed' in str(p.relative_to(sow_dir)).lower():
            continue
        rel   = p.relative_to(BASE_DIR)
        parts = list(rel.parts)
        category = " / ".join(parts[2:-1]) if len(parts) > 3 else (parts[2] if len(parts) > 2 else "")
        docs.append({
            "path":     str(p),
            "rel_path": str(rel).replace("\\", "/"),
            "name":     p.name,
            "category": category,
            "ext":      p.suffix.lower(),
            "size_kb":  round(p.stat().st_size / 1024),
        })
    return docs


def suggest_docs(query: str):
    """Keyword-match user query against SOW doc names."""
    docs  = list_sow_docs()
    query = query.strip()
    if not query:
        return docs
    terms = re.split(r'[\s,]+', query.lower())
    scored = []
    for doc in docs:
        haystack = f"{doc['name']} {doc['category']}".lower()
        score = sum(2 if term in doc['name'].lower() else 1
                    for term in terms if term and term in haystack)
        if score > 0:
            scored.append({**doc, "relevance": score})
    scored.sort(key=lambda x: -x["relevance"])
    return scored if scored else docs


# ─── Initial dashboard analysis ──────────────────────────────────────────────

def generate_initial_analysis(user_name: str = "דנה") -> dict:
    """Return a simple static welcome message — no AI call needed."""
    docs = list_sow_docs()
    return {
        "greeting": f"שלום {user_name}! ברוכה הבאה למערכת רכש חכם.",
        "doc_count": len(docs),
    }


# ─── Chat-based document recommendation ──────────────────────────────────────

def _detect_analytical_request(message: str) -> bool:
    """Detect if message is a direct analytical/pipeline request vs a chat request."""
    # Hebrew analytical keywords and stems
    analytical_keywords = [
        "ניתוח", "analysis", "analyze",
        "עדכ", "update",  # עדכן, עדכון, עדכני, etc.
        "יצ", "create", "generate",  # יצור, יצירת, יצירה, etc.
        "השו", "compare", "comparison",  # השוו, השוואה, etc.
        "סקור", "review", "examine",
        "דריש", "requirements", "specification",  # דרישות, דרישה, etc.
        "2026", "2027", "2028", "2029", "2030", "202", # years
        "תכול", "sow", "statement of work",  # תכולה, תכולת, etc.
        "מסמ", "document", "report",  # מסמך, מסמכים, etc.
        "טיוט", "draft",  # טיוטה, טיוטות, etc.
        "דלק", "fuel",  # דלק
        "סליק", "payment",  # סליקה, סליקת, etc.
        "גיבוש", "consolidate",  # גיבוש
    ]
    msg_lower = message.lower()

    # Check if any keyword matches
    match = any(kw in msg_lower for kw in analytical_keywords)

    # Also auto-trigger for long requests (likely analytical)
    is_long = len(message) > 80

    return match or is_long


def _fallback_recommend(message: str, available_docs: list) -> list:
    """Pick relevant docs by keyword when Gemini is unavailable."""
    msg_lower = message.lower()
    if "דלק" in message or "fuel" in msg_lower:
        docs = [d for d in available_docs if 'דלק' in d.get('name', '') or 'fuel' in d.get('name', '').lower()]
    elif "ניקוי" in message or "clean" in msg_lower:
        docs = [d for d in available_docs if 'ניקוי' in d.get('name', '')]
    else:
        docs = [d for d in available_docs if 'sow' in d.get('name', '').lower()]
    return docs[:5] if docs else available_docs[:4]


def chat_recommend(message: str, history: list, available_docs: list, selected_files: list | None = None) -> dict:
    """Use Gemini Flash to recommend docs and suggest actions, or trigger direct pipeline."""

    selected_files = selected_files or []

    # ── Priority path: user already selected files → go straight to pipeline ──
    if selected_files:
        # Build recommended list from user-selected files (match by path)
        path_set = set(selected_files)
        user_docs = [d for d in available_docs if d.get('path') in path_set]
        # Fall back to the raw paths if not found in available_docs
        if not user_docs:
            user_docs = [{"path": p, "name": Path(p).name} for p in selected_files]

        doc_names = ", ".join([d['name'][:40] for d in user_docs[:5]])
        return {
            "message": f"✓ מתחיל עיבוד\n📄 מסמכים נבחרים: {doc_names}\n⏳ טוען סוכנים (מומחה תחום, משפטי, שוק)...",
            "recommended_files": user_docs,
            "suggested_actions": [],
            "proposed_prompt": message,
            "ready_to_generate": True,
            "auto_start": True,
        }

    # ── Analytical request detected → auto-recommend files and start pipeline ──
    if _detect_analytical_request(message):
        # Recommend relevant SOW documents for analysis
        msg_lower = message.lower()

        # First, try to find documents matching the specific domain mentioned in the message
        domain_matched_docs = []
        if "דלק" in message or "fuel" in msg_lower:
            domain_matched_docs = [d for d in available_docs if 'דלק' in d.get('name', '') or 'fuel' in d.get('name', '').lower() or d.get('category', '').lower() in ['דלק', 'fuel']]
        elif "ניקוי" in message or "clean" in msg_lower:
            domain_matched_docs = [d for d in available_docs if 'ניקוי' in d.get('name', '') or 'clean' in d.get('name', '').lower()]

        # Then fall back to SOW documents in general
        if domain_matched_docs:
            recommended = domain_matched_docs[:5]
        else:
            sow_docs = [d for d in available_docs if 'sow' in d.get('name', '').lower() or 'sow' in d.get('category', '').lower()]
            recommended = (sow_docs[:4] if sow_docs else available_docs[:4])

        doc_names = ", ".join([d['name'][:40] for d in recommended]) if recommended else "כל המסמכים הזמינים"

        return {
            "message": f"✓ מתחיל ניתוח מעמיק\n📄 מסמכים: {doc_names}\n⏳ טוען סוכנים (מומחה תחום, משפטי, שוק)...",
            "recommended_files": recommended,
            "suggested_actions": [],
            "proposed_prompt": message,  # Use message as context
            "ready_to_generate": True,
            "auto_start": True,  # Tell frontend to skip prompt confirmation and start immediately
        }

    doc_list = "\n".join(
        f"- {d['name']} (קטגוריה: {d['category']}, {d['ext']}, {d['size_kb']}KB)"
        for d in available_docs
    )

    hist_text = ""
    for h in history[-10:]:
        role = "משתמש" if h.get("role") == "user" else "עוזר"
        hist_text += f"{role}: {h.get('content', '')}\n"

    prompt = f"""אתה עוזר AI לרכש של משרד הביטחון. אתה עוזר למשתמש לבחור מסמכי SOW ולתכנן פעולות.

{IDF_SYSTEM_CONTEXT}

מסמכים זמינים במערכת:
{doc_list}

היסטוריית שיחה:
{hist_text}

הודעת המשתמש: {message}

ענה בעברית. החזר JSON בלבד:
{{
  "message": "תשובה למשתמש — בעברית, תמציתית",
  "recommended_files": [
    {{"name": "שם הקובץ בדיוק כמו ברשימה", "reason": "למה רלוונטי"}}
  ],
  "suggested_actions": [
    {{"id": "action_1", "label": "תווית קצרה לפעולה", "description": "תיאור מלא של הפעולה", "type": "sow_update|research|analysis|comparison"}}
  ],
  "proposed_prompt": null,
  "ready_to_generate": false
}}

כללים:
- כשהמשתמש בוחר קבצים: הצע 2-3 פעולות ב-suggested_actions שהמשתמש יכול לבחור
- suggested_actions: פעולות מומלצות (עדכון SOW, מחקר שוק, ניתוח סיכונים, השוואת מסמכים)
- אם המשתמש מאשר פעולה: הצע prompt מפורט ב-proposed_prompt
- אם המשתמש מאשר prompt: סמן ready_to_generate=true
- שם הקובץ חייב להיות זהה לרשימת המסמכים
- היה ספציפי לתחום — אל תיתן המלצות גנריות"""

    try:
        resp = _client.models.generate_content(model=FLASH_MODEL, contents=prompt)
        data = _parse_json_robust(resp.text)
        name_to_doc = {d['name']: d for d in available_docs}
        for rf in data.get('recommended_files', []):
            doc = name_to_doc.get(rf.get('name'))
            if doc:
                rf['path'] = doc['path']
                rf['category'] = doc.get('category', '')
                rf['size_kb'] = doc.get('size_kb', 0)
        return data
    except (json.JSONDecodeError, Exception) as e:
        print(f"Error in chat_recommend: {type(e).__name__}: {str(e)[:200]}")
        # If user had files selected, use those and go to pipeline
        if selected_files:
            path_set = set(selected_files)
            user_docs = [d for d in available_docs if d.get('path') in path_set]
            if not user_docs:
                user_docs = [{"path": p, "name": Path(p).name} for p in selected_files]
            return {
                "message": f"✓ מתחיל עיבוד עם {len(user_docs)} מסמכים שנבחרו...",
                "recommended_files": user_docs,
                "suggested_actions": [],
                "proposed_prompt": message,
                "ready_to_generate": True,
                "auto_start": True,
            }
        # Otherwise try to recommend docs and start pipeline
        if _detect_analytical_request(message):
            recommended = _fallback_recommend(message, available_docs)
            return {
                "message": f"✓ מעבד את הבקשה — {len(recommended)} מסמכים נבחרו אוטומטית...",
                "recommended_files": recommended,
                "suggested_actions": [],
                "proposed_prompt": message,
                "ready_to_generate": True,
                "auto_start": True,
            }
        return {
            "message": "בחרי מסמכים מהסרגל הצדדי ושלחי שוב.",
            "recommended_files": [],
            "suggested_actions": [],
            "proposed_prompt": None,
            "ready_to_generate": False,
        }


# ─── File upload helpers ──────────────────────────────────────────────────────

def _extract_docx_text(path: Path) -> str:
    doc_obj = docx.Document(str(path))
    paragraphs = [p.text for p in doc_obj.paragraphs if p.text.strip()]
    for table in doc_obj.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(' | '.join(cells))
    return '\n'.join(paragraphs)


def _ascii_display_name(filename: str) -> str:
    norm = unicodedata.normalize('NFKD', filename)
    safe = ''.join(c if ord(c) < 128 else '_' for c in norm)
    safe = re.sub(r'_+', '_', safe).strip('_')
    ext  = Path(filename).suffix
    return safe if safe.replace('_', '') else f"document{ext}"


def _upload_files(file_paths: list) -> list:
    uploaded = []
    for i, path in enumerate(file_paths):
        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"File not found: {path}")
        disp_name = _ascii_display_name(p.name)
        ext = p.suffix.lower()
        if ext == '.pdf':
            mime = 'application/pdf'
            cfg = UploadFileConfig(mime_type=mime, display_name=disp_name)
            with open(p, 'rb') as fh:
                f = _client.files.upload(file=fh, config=cfg)
        elif ext in ('.docx',):
            text = _extract_docx_text(p)
            cfg = UploadFileConfig(mime_type='text/plain', display_name=disp_name)
            f = _client.files.upload(file=io.BytesIO(text.encode('utf-8')), config=cfg)
        elif ext in ('.doc',):
            cfg = UploadFileConfig(mime_type='text/plain', display_name=disp_name)
            with open(p, 'rb') as fh:
                f = _client.files.upload(file=fh, config=cfg)
        else:
            cfg = UploadFileConfig(mime_type='text/plain', display_name=disp_name)
            with open(p, 'rb') as fh:
                f = _client.files.upload(file=fh, config=cfg)
        for _ in range(60):
            if f.state != FileState.PROCESSING:
                break
            time.sleep(2)
            f = _client.files.get(name=f.name)
        if f.state != FileState.ACTIVE:
            raise RuntimeError(f"File '{p.name}' failed to activate (state={f.state})")
        uploaded.append(f)
    return uploaded


def _clean_json(raw: str) -> str:
    raw = raw.strip()
    if raw.startswith("```"):
        lines = raw.split("\n")
        end   = len(lines) - 1 if lines[-1].strip() == "```" else len(lines)
        raw   = "\n".join(lines[1:end])
    return raw.strip()


def _parse_json_robust(raw: str) -> dict:
    """Parse JSON from Gemini with repair for common LLM issues."""
    cleaned = _clean_json(raw)
    # Try direct parse
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Apply multiple repair strategies
    def _repair(s):
        # Fix trailing commas: ,} or ,]
        s = re.sub(r',\s*([}\]])', r'\1', s)
        # Fix missing commas between values: "value"\n"key" or "value" "key"
        s = re.sub(r'"\s*\n\s*"', '",\n"', s)
        # Fix missing commas after numbers: 123\n"key"
        s = re.sub(r'(\d)\s*\n\s*"', r'\1,\n"', s)
        # Fix missing commas after booleans: true\n"key" or false\n"key"
        s = re.sub(r'(true|false|null)\s*\n\s*"', r'\1,\n"', s)
        # Fix missing commas after closing brackets: }\n"key" or ]\n"key"
        s = re.sub(r'([}\]])\s*\n\s*"', r'\1,\n"', s)
        # Fix missing commas after closing brackets before opening: }{ or ][
        s = re.sub(r'([}\]])\s*\n\s*([{\[])', r'\1,\n\2', s)
        return s

    repaired = _repair(cleaned)
    try:
        return json.loads(repaired)
    except json.JSONDecodeError:
        pass

    # Try extracting the JSON object/array
    for start_char, end_char in [('{', '}'), ('[', ']')]:
        start = cleaned.find(start_char)
        end = cleaned.rfind(end_char)
        if start != -1 and end != -1 and end > start:
            subset = cleaned[start:end + 1]
            try:
                return json.loads(_repair(subset))
            except json.JSONDecodeError:
                pass

    # Last resort: raise
    return json.loads(cleaned)


# ─── Quality review agent ────────────────────────────────────────────────────

def _quality_review(result: dict, domain_ctx: dict) -> dict:
    """Agent 4: Review all output for domain relevance and accuracy."""
    review_prompt = f"""אתה מומחה בקרת איכות לרכש ממשלתי ישראלי.
בדוק שכל הפלט הבא רלוונטי ומדויק לתחום: {domain_ctx.get('domain', 'כללי')}

{IDF_SYSTEM_CONTEXT}

תחום: {domain_ctx.get('domain', 'כללי')}
רגולציה רלוונטית: {domain_ctx.get('regulations', '')}
מומחיות נדרשת: {domain_ctx.get('expertise', '')}
ספקים רלוונטיים: {domain_ctx.get('vendors_focus', '')}
מה להימנע ממנו: {domain_ctx.get('avoid', '')}

הפלט לבדיקה:
{json.dumps(result, ensure_ascii=False)[:8000]}

בדוק:
1. האם כל ההמלצות רלוונטיות לתחום הספציפי?
2. האם יש המלצות גנריות שצריך להסיר או לשנות?
3. האם הספקים המומלצים באמת עוסקים בתחום?
4. האם הרגולציה שצוינה נכונה?
5. האם יש תוכן שלא קשור לתחום שצריך להסיר?

החזר JSON:
{{
  "is_relevant": true/false,
  "issues_found": ["בעיה 1", "בעיה 2"],
  "cleaned_vendors": [/* ספקים שעברו סינון — רק רלוונטיים */],
  "notes": "הערות כלליות"
}}"""

    try:
        resp = _client.models.generate_content(model=FLASH_MODEL, contents=review_prompt)
        review = _parse_json_robust(resp.text)
        # If vendors were cleaned, update them
        if review.get("cleaned_vendors") and result.get("market_analysis", {}).get("vendors"):
            clean_names = {v.get("name") for v in review["cleaned_vendors"]}
            if clean_names:
                orig_vendors = result["market_analysis"]["vendors"]
                result["market_analysis"]["vendors"] = [
                    v for v in orig_vendors
                    if v.get("name") in clean_names or v.get("status") == "qualified"
                ]
        return result
    except Exception:
        return result


# ─── Markdown export ─────────────────────────────────────────────────────────

def generate_markdown_output(result: dict) -> str:
    """Generate IDF-styled markdown with TOC and enhanced formatting."""
    title = result.get("tender_title", "SOW מעודכן")
    domain = result.get("domain", "")
    sections = []
    for doc in result.get("documents", []):
        if doc.get("type") == "sow":
            sections = doc.get("sections", [])
            break

    accepted = [s for s in sections if s.get("status") != "reject"]
    today = datetime.now().strftime('%d/%m/%Y')

    lines = []

    # ── Cover page with metadata ────────────────────────────────────────────
    lines += [
        "---",
        "מסמך: **סודי / RESTRICTED** — לשימוש פנימי בלבד",
        "---",
        "",
        f"# {title}",
        "",
        "**מסמך תכולת עבודה (Statement of Work)**",
        "",
        f"| שדה | ערך |",
        f"|-----|-----|",
        f"| גוף מזמין | מדינת ישראל — משרד הביטחון |",
        f"| תחום | {domain} |",
        f"| גרסה | 2026 — מעודכן ע\"י מערכת AI לרכש חכם |",
        f"| תאריך | {today} |",
        f"| סיווג | מוגבל — לשימוש פנימי בלבד |",
        "",
        "> מסמך זה וכל נספחיו הינו בבעלות בלעדית של משרד הביטחון.",
        "> חל איסור להעבירו לצד שלישי ללא אישור בכתב ומראש של משרד הביטחון.",
        "",
        "---",
        "",
    ]

    # ── Table of Contents ──────────────────────────────────────────────────
    lines.append("## תוכן עניינים")
    lines.append("")
    for s in accepted:
        num = s.get("number", "")
        sec_title = s.get("title", "")
        ct = s.get("change_type", "unchanged")
        badge = " *(חדש)*" if ct == "added" else " *(עודכן)*" if ct == "modified" else ""
        anchor = re.sub(r'[^\w\s-]', '', f"{num} {sec_title}", flags=re.UNICODE).strip().lower()
        anchor = re.sub(r'[\s]+', '-', anchor)
        lines.append(f"- [{num}. {sec_title}{badge}](#{anchor})")
    lines.append("")
    lines.append("---")
    lines.append("")

    # ── Sections ────────────────────────────────────────────────────────────
    for s in accepted:
        num = s.get("number", "")
        sec_title = s.get("title", "")
        ct = s.get("change_type", "unchanged")
        reason = s.get("change_reason", "")
        risk = s.get("risk_level", "low")
        text = s.get("updated") or s.get("original") or ""

        lines.append(f"## {num}. {sec_title}")
        lines.append("")

        if ct == "modified":
            priority = "P1 — שינוי קריטי" if risk == "high" else "P2 — שיפור חשוב"
            lines.append(f"> **[{priority}]**")
            lines.append(f"> *סיבה לשינוי: {reason}*")
            lines.append("")
        elif ct == "added":
            lines.append(f"> **[סעיף חדש — P2]**")
            lines.append(f"> *סיבה: {reason}*")
            lines.append("")

        # Detect and render pipe-separated table rows
        lines += _text_to_md(text)
        lines.append("")
        lines.append("---")
        lines.append("")

    # ── Q&A Appendix ───────────────────────────────────────────────────────
    qa_items = []
    for doc in result.get("documents", []):
        if doc.get("type") == "vendor_qa":
            qa_items = doc.get("qa_items", [])
            break

    if qa_items:
        lines.append("## נספח א׳ — שאלות הבהרה ותשובות לספקים")
        lines.append("")
        lines.append("| # | שאלה | תשובה | קטגוריה | השפעה |")
        lines.append("|---|------|--------|----------|-------|")
        for i, q in enumerate(qa_items, 1):
            if q.get("status") != "reject":
                q_text = q.get('question', '').replace('|', '\\|').replace('\n', ' ')
                a_text = q.get('answer', '').replace('|', '\\|').replace('\n', ' ')
                cat    = q.get('category', '').replace('|', '\\|')
                impact = q.get('vendor_impact', '').replace('|', '\\|')
                lines.append(f"| {i} | {q_text} | {a_text} | {cat} | {impact} |")
        lines.append("")

    lines.append("---")
    lines.append(f"*מסמך זה הופק אוטומטית ע\"י מערכת AI לרכש חכם — {today}*")
    lines.append("")
    lines.append("*מיועד לעיון לפני אישור סופי ע\"י הגורמים המוסמכים במשרד הביטחון*")

    return "\n".join(lines)


def _text_to_md(text: str) -> list:
    """Convert plain text to markdown lines, detecting embedded tables."""
    output_lines = []
    prev_was_table = False
    for para in text.split('\n'):
        stripped = para.strip()
        if not stripped:
            output_lines.append("")
            prev_was_table = False
            continue
        # Detect pipe-table rows: "col1 | col2 | col3"
        if ' | ' in stripped and stripped.count(' | ') >= 1:
            if not prev_was_table:
                output_lines.append(stripped)
                cols = stripped.count(' | ') + 1
                output_lines.append('|'.join(['---'] * cols))
                prev_was_table = True
            else:
                output_lines.append(stripped)
        else:
            output_lines.append(stripped)
            prev_was_table = False
    return output_lines


# ─── Main pipeline ────────────────────────────────────────────────────────────

def run_pipeline(file_paths: list, user_context: str, sub_type: str):
    """Full multi-agent Gemini pipeline with domain awareness and quality review."""

    # Detect domain
    domain_ctx = _detect_domain(file_paths, user_context)
    domain_name = domain_ctx.get("domain", "כללי")

    yield {"agent": "orchestrator", "status": "running",
           "msg": f"מזהה תחום: {domain_name} — מנתב לסוכנים מתאימים..."}

    # Step 1: Upload files
    n = len(file_paths)
    yield {"agent": "domain_expert", "status": "running",
           "msg": f"מעלה {n} קבצים ל-Gemini..."}

    try:
        gemini_files = _upload_files(file_paths)
    except Exception as e:
        yield {"agent": "domain_expert", "status": "error",
               "msg": f"שגיאת העלאה: {str(e)[:200]}"}
        yield {"type": "error", "msg": str(e)}
        return

    names = [Path(f_).name for f_ in file_paths]
    yield {"agent": "domain_expert", "status": "running",
           "msg": f"הועלו {len(gemini_files)} קבצים: {', '.join(names[:3])}{'...' if len(names) > 3 else ''}"}

    yield {"agent": "orchestrator", "status": "running",
           "msg": f"ניתוב לסוכן מומחה {domain_name} — מתחיל ניתוח מסמכים"}

    # Agent 1: Domain Expert
    yield {"agent": "domain_expert", "status": "running",
           "msg": "קורא ומנתח את כל המסמכים..."}

    # Estimate document length for target
    total_chars = sum(Path(fp).stat().st_size for fp in file_paths if Path(fp).exists())
    estimated_pages = max(10, total_chars // 2000)

    domain_prompt = f"""You are an expert procurement analyst for Israeli defense procurement (רכש ממשלתי ישראלי).

{IDF_SYSTEM_CONTEXT}

DOMAIN SPECIALIZATION: {domain_name}
Relevant regulations: {domain_ctx.get('regulations', '')}
Required expertise: {domain_ctx.get('expertise', '')}
IMPORTANT - AVOID: {domain_ctx.get('avoid', '')}

Read ALL attached documents carefully. These are SOW (מסמך תכולת עבודה) documents in Hebrew.

User request: {user_context or "סקור ועדכן את המסמכים לשנת 2026"}
Tender type: {sub_type or domain_name}

CRITICAL DOCUMENT REQUIREMENTS:
Before writing anything: analyze the original document structure.
Count pages, identify tables, TOC, section hierarchy, and section count.
Your new document MUST be at least {estimated_pages} pages of Hebrew content.

Tasks:
1. UPDATED SOW — Process EVERY section in the original. Do NOT skip or summarize any section.
   - Reproduce each original section fully, then expand and update it for 2026.
   - New gap sections: minimum 200 words each.
   - ALL tables from the original must appear as structured data in the updated version.
   - Preserve TOC structure, section numbering, and subsection hierarchy.
   - Every "updated" field must be AT LEAST as long as its "original" — never shorter.
   - Be MORE detailed and sharper than the original — specific criteria, exact amounts, article numbers.
   - Do NOT use "..." or truncate any content.
   - change_reason must be specific and actionable (not generic).
   - Target: minimum {estimated_pages} pages worth of content across all sections.

2. VENDOR Q&A — Min 8 Q&A pairs. Each answer minimum 60 words. Domain-specific to {domain_name} only.

CRITICAL: A document quality check will run after. Any section where updated text is shorter
than original text will be flagged as a failure. Write complete, comprehensive content.

All recommendations, vendors, regulations MUST be specific to {domain_name}.
Do NOT recommend vendors from unrelated industries.
Do NOT provide generic recommendations.
Be specific, be accurate, be relevant.

Follow this IDF SOW structure when applicable:
1. כללי — מטרת ההתקשרות
2. מסמכים ישימים + תקנים
3. הגדרות
4. LOTs — פיצול לעידוד תחרות
5. חבילת השירות (זמני אספקה, סכום מינימום, עיצומים)
6. מאפיינים טכניים ותקנים
7. הבטחת איכות
8. Recall / הוצאה מהשוק
9. שרשרת קירור (אם רלוונטי)
10. אנשי קשר
11. ביטחון מידע
12. מצב חירום
13. פיקוח ועיצומים
14. נספחים

Return ONLY valid JSON:
{{
  "tender_title": "Title — Hebrew",
  "sow_sections": [
    {{
      "id": "sec_1", "number": "1", "title": "Section title — Hebrew",
      "original": "Verbatim from doc (null if new)",
      "updated": "Full updated text — detailed, specific — Hebrew",
      "change_type": "modified|added|unchanged",
      "change_reason": "Why — Hebrew",
      "agent": "domain_expert"
    }}
  ],
  "qa_items": [
    {{
      "id": "qa_1", "question": "Question", "answer": "Answer — Hebrew",
      "category": "טכני|משפטי|מסחרי|תפעולי",
      "vendor_impact": "high|medium|low", "source": "Source"
    }}
  ]
}}

All text MUST be in Hebrew. Be specific to {domain_name}."""

    try:
        yield {"agent": "domain_expert", "status": "running",
               "msg": f"שולח ל-Gemini — ניתוח מעמיק של תחום {domain_name}..."}
        resp1        = _client.models.generate_content(model=MODEL,
                         contents=[*gemini_files, domain_prompt],
                         config={"response_mime_type": "application/json"})
        data1        = _parse_json_robust(resp1.text)
        sow_sections = data1.get("sow_sections", [])
        qa_items     = data1.get("qa_items", [])
        tender_title = data1.get("tender_title", f"SOW — {sub_type} 2026")
    except json.JSONDecodeError as e:
        # Retry once
        try:
            yield {"agent": "domain_expert", "status": "running",
                   "msg": "ניסיון חוזר — שגיאת JSON, מנסה שוב..."}
            resp1        = _client.models.generate_content(model=MODEL,
                             contents=[*gemini_files, domain_prompt],
                             config={"response_mime_type": "application/json"})
            data1        = _parse_json_robust(resp1.text)
            sow_sections = data1.get("sow_sections", [])
            qa_items     = data1.get("qa_items", [])
            tender_title = data1.get("tender_title", f"SOW — {sub_type} 2026")
        except Exception as e2:
            yield {"agent": "domain_expert", "status": "error",
                   "msg": f"שגיאת JSON מ-Gemini: {str(e2)[:120]}"}
            yield {"type": "error", "msg": f"JSON parse error from Agent 1: {e2}"}
            return
    except Exception as e:
        yield {"agent": "domain_expert", "status": "error",
               "msg": f"שגיאת Gemini: {str(e)[:200]}"}
        yield {"type": "error", "msg": str(e)}
        return

    yield {"agent": "domain_expert", "status": "done",
           "msg": f"הושלם — {len(sow_sections)} סעיפי SOW, {len(qa_items)} שאלות"}

    # Agent 2: Legal + Market
    yield {"agent": "legal", "status": "running",
           "msg": f"בוחן תנאי סף, סיכונים משפטיים ושוק ספקים ל{domain_name}..."}

    legal_prompt = f"""You are a legal advisor and market analyst for Israeli government procurement.

{IDF_SYSTEM_CONTEXT}

DOMAIN: {domain_name}
Regulations: {domain_ctx.get('regulations', '')}
Vendors focus: {domain_ctx.get('vendors_focus', '')}
CRITICAL - AVOID: {domain_ctx.get('avoid', '')}

User context: {user_context or "ניתוח מכרז"}
Tender type: {sub_type or domain_name}

SOW sections from domain expert:
{json.dumps({"sections": sow_sections}, ensure_ascii=False, indent=2)[:6000]}

Tasks:
1. Add legal risk assessment to every section
2. Create market analysis with REAL vendors specific to {domain_name}

CRITICAL:
- Only recommend vendors that ACTUALLY work in {domain_name}
- Do NOT recommend tech companies for fuel, do NOT recommend food companies for furniture
- {domain_ctx.get('avoid', '')}
- Be specific — use real Israeli vendor names when possible

Return ONLY valid JSON:
{{
  "sow_sections": [/* All sections with: risk_level, threshold_difficulty, threshold_note */],
  "market_analysis": {{
    "summary": "Market overview — Hebrew, specific to {domain_name}",
    "vendors": [
      {{
        "name": "Real vendor name", "country": "Country", "category": "{domain_name}",
        "fit_score": 4, "strengths": "Hebrew", "barriers": "Hebrew",
        "status": "qualified|conditional|disqualified"
      }}
    ],
    "challenges": [
      {{
        "id": "ch_1", "title": "Hebrew", "severity": "critical|high|medium|low",
        "must_have": true, "category": "technical|legal|operational|security|market",
        "description": "Hebrew", "affected_vendor_count": 5, "recommendation": "Hebrew"
      }}
    ]
  }},
  "threshold_summary": {{"easy": 0, "medium": 0, "hard": 0}}
}}

Min 5 vendors (specific to {domain_name}!), 6 challenges. All Hebrew."""

    try:
        yield {"agent": "legal", "status": "running",
               "msg": "ניתוח משפטי ושוק ספקים..."}
        resp2 = _client.models.generate_content(model=MODEL, contents=legal_prompt,
                    config={"response_mime_type": "application/json"})
        data2 = _parse_json_robust(resp2.text)
        sow_sections      = data2.get("sow_sections", sow_sections)
        market_analysis   = data2.get("market_analysis",
                                       {"summary": "", "vendors": [], "challenges": []})
        threshold_summary = data2.get("threshold_summary",
                                       _count_thresholds(sow_sections))
    except Exception as e:
        yield {"agent": "legal", "status": "warning",
               "msg": f"ניתוח חלקי ({str(e)[:80]}) — ממשיך עם תוצאות חלקיות"}
        market_analysis   = {"summary": "", "vendors": [], "challenges": []}
        threshold_summary = _count_thresholds(sow_sections)

    yield {"agent": "legal", "status": "done",
           "msg": f"הושלם — {len(market_analysis.get('vendors', []))} ספקים, {len(market_analysis.get('challenges', []))} אתגרים"}

    # Agent 3: Orchestrator — assemble
    yield {"agent": "orchestrator", "status": "running",
           "msg": "מרכב מסמכים ובודק עקביות..."}

    for s in sow_sections:
        s.setdefault("status", "pending")
        if not s.get("updated"):
            s["updated"] = s.get("original") or ""
    for q in qa_items:
        q.setdefault("status", "pending")

    docs = [{
        "id": "doc_sow",
        "type": "sow",
        "title": tender_title,
        "description": (
            f"{sum(1 for s in sow_sections if s.get('change_type') != 'unchanged')} שינויים | "
            f"{sum(1 for s in sow_sections if s.get('change_type') == 'added')} סעיפים חדשים"
        ),
        "sections": sow_sections,
    }]
    if qa_items:
        docs.append({
            "id": "doc_qa",
            "type": "vendor_qa",
            "title": "תשובות לשאלות הבהרה — ספקים",
            "description": f"{len(qa_items)} שאלות | תשובות מחייבות",
            "qa_items": qa_items,
        })

    result = {
        "request_summary": (user_context or f"עדכון SOW ל{sub_type} 2026")[:300],
        "tender_title":     tender_title,
        "sub_type":         sub_type,
        "domain":           domain_name,
        "documents":        docs,
        "market_analysis":  market_analysis,
        "threshold_summary": threshold_summary,
    }

    # Agent 4: Quality Review
    yield {"agent": "orchestrator", "status": "running",
           "msg": f"סוכן בקרת איכות בודק רלוונטיות לתחום {domain_name}..."}

    result = _quality_review(result, domain_ctx)

    yield {"agent": "orchestrator", "status": "done",
           "msg": f"הושלם — {len(docs)} מסמכים, {len(sow_sections)} סעיפים"}

    yield {"type": "complete", "result": result}


def _count_thresholds(sections: list) -> dict:
    c = {"easy": 0, "medium": 0, "hard": 0}
    for s in sections:
        d = s.get("threshold_difficulty")
        if d in c:
            c[d] += 1
    return c
