#!/usr/bin/env python3
"""
create_gdocs.py — Create Google Docs from SOW fixed markdown files.

Steps:
  1. Converts each fixed/*.md file to a DOCX (RTL Hebrew, styled like original MOD docs).
  2. Uploads each DOCX to Google Drive.
  3. Converts to Google Docs format on Drive.
  4. Saves public sharing URLs to gdoc_links.json.

Requirements:
  pip install python-docx google-api-python-client google-auth-httplib2 google-auth-oauthlib

Google OAuth setup:
  1. Go to https://console.cloud.google.com/
  2. Create a project → Enable "Google Drive API"
  3. Create OAuth credentials (Desktop app) → Download as credentials.json
  4. Place credentials.json in the same folder as this script (data/SOW/)

Usage:
  python create_gdocs.py

Output:
  gdoc_links.json  — maps original filename to Google Doc URL
  docx/            — generated DOCX files (kept for reference)
"""

import json
import os
import sys
import re
from pathlib import Path

# ─── Dependency check ────────────────────────────────────────────────────────
def install_deps():
    import subprocess
    packages = [
        "google-api-python-client",
        "google-auth-httplib2",
        "google-auth-oauthlib",
        "python-docx"
    ]
    print("Installing dependencies...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "--quiet"] + packages)
    print("Dependencies installed.\n")

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload
except ImportError:
    install_deps()
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaFileUpload


# ─── Paths ────────────────────────────────────────────────────────────────────
BASE_DIR    = Path(__file__).parent
FIXED_DIR   = BASE_DIR / "fixed"
DOCX_DIR    = BASE_DIR / "docx"
CREDS_FILE  = BASE_DIR / "credentials.json"
TOKEN_FILE  = BASE_DIR / "token.json"
LINKS_FILE  = BASE_DIR / "gdoc_links.json"

SCOPES = ["https://www.googleapis.com/auth/drive"]

DRIVE_FOLDER_NAME = "SOW Analysis — מסמכים משופרים"


# ─── Markdown → DOCX ─────────────────────────────────────────────────────────
def set_rtl(paragraph):
    """Set RTL direction on a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def make_run_rtl(run):
    """Mark run text as RTL Hebrew."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)


def add_rtl_paragraph(doc, text, style_name="Normal", bold=False, size=None, color=None):
    """Add an RTL paragraph to the document."""
    para = doc.add_paragraph(style=style_name)
    set_rtl(para)
    run = para.add_run(text)
    make_run_rtl(run)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para


def md_to_docx(md_path: Path, docx_path: Path):
    """Convert a markdown file to a styled DOCX with RTL Hebrew."""
    doc = Document()

    # Page setup: A4, RTL
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.right_margin = Cm(2.5)
    section.left_margin  = Cm(2.5)
    section.top_margin   = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # Document-level RTL
    doc.core_properties.language = "he-IL"
    doc_body = doc.element.find(qn("w:body"))
    sectPr = doc_body.find(qn("w:sectPr"))
    if sectPr is None:
        sectPr = OxmlElement("w:sectPr")
        doc_body.append(sectPr)
    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)

    # Styles
    styles = doc.styles
    try:
        normal = styles["Normal"]
        normal.font.name = "Arial"
        normal.font.size = Pt(11)
    except Exception:
        pass

    lines = md_path.read_text(encoding="utf-8").splitlines()
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            in_table = False
            return
        # Filter separator rows (|---|---|)
        data_rows = [r for r in table_rows if not re.match(r"^\|[-| :]+\|?\s*$", r)]
        if not data_rows:
            in_table = False
            table_rows = []
            return
        # Parse cells
        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in row.strip("|").split("|")]
            parsed.append(cells)
        if not parsed:
            in_table = False
            table_rows = []
            return
        col_count = max(len(r) for r in parsed)
        table = doc.add_table(rows=len(parsed), cols=col_count)
        table.style = "Table Grid"
        for r_idx, row in enumerate(parsed):
            for c_idx, cell_text in enumerate(row):
                cell = table.rows[r_idx].cells[c_idx]
                cell.text = cell_text
                para = cell.paragraphs[0]
                set_rtl(para)
                for run in para.runs:
                    make_run_rtl(run)
                    if r_idx == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                if r_idx == 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "1E3A5F")
                    shd.set(qn("w:val"), "clear")
                    tcPr.append(shd)
        in_table = False
        table_rows = []
        doc.add_paragraph()  # spacing after table

    for line in lines:
        stripped = line.strip()

        # Detect table
        if stripped.startswith("|") and stripped.endswith("|"):
            in_table = True
            table_rows.append(stripped)
            continue
        elif in_table:
            flush_table()

        if not stripped:
            doc.add_paragraph()
            continue

        # Blockquote (P1/P2/P3 annotations)
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            para = doc.add_paragraph()
            set_rtl(para)
            para.paragraph_format.left_indent  = Cm(0.5)
            para.paragraph_format.right_indent = Cm(0.5)
            run = para.add_run(text)
            make_run_rtl(run)
            run.font.color.rgb = RGBColor(92, 64, 5)
            run.italic = True
            # Yellow background via shading
            pPr = para._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "FFFBEB")
            shd.set(qn("w:val"), "clear")
            pPr.append(shd)
            continue

        # Headings
        if stripped.startswith("### "):
            para = doc.add_paragraph(stripped[4:], style="Heading 3")
            set_rtl(para)
        elif stripped.startswith("## "):
            para = doc.add_paragraph(stripped[3:], style="Heading 2")
            set_rtl(para)
        elif stripped.startswith("# "):
            para = doc.add_paragraph(stripped[2:], style="Heading 1")
            set_rtl(para)
            for run in para.runs:
                run.font.color.rgb = RGBColor(30, 58, 95)
        # Horizontal rule
        elif stripped.startswith("---"):
            para = doc.add_paragraph()
            set_rtl(para)
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:color"), "AAAAAA")
            pBdr.append(bottom)
            pPr.append(pBdr)
        # List items
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(stripped[2:], style="List Bullet")
            set_rtl(para)
            for run in para.runs:
                make_run_rtl(run)
        # Numbered list
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            para = doc.add_paragraph(text, style="List Number")
            set_rtl(para)
            for run in para.runs:
                make_run_rtl(run)
        # Normal paragraph (handle **bold**)
        else:
            para = doc.add_paragraph()
            set_rtl(para)
            # Simple bold parsing
            parts = re.split(r"\*\*([^*]+)\*\*", stripped)
            for i, part in enumerate(parts):
                if not part:
                    continue
                run = para.add_run(part)
                make_run_rtl(run)
                if i % 2 == 1:  # odd = inside **...**
                    run.bold = True

    if in_table:
        flush_table()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    return docx_path


# ─── Google Drive ─────────────────────────────────────────────────────────────
def get_drive_service():
    """Authenticate and return a Google Drive service object."""
    if not CREDS_FILE.exists():
        print(f"\n❌ credentials.json not found at: {CREDS_FILE}")
        print("Steps to fix:")
        print("  1. Go to https://console.cloud.google.com/")
        print("  2. Create/select a project → APIs & Services → Enable APIs → Google Drive API")
        print("  3. Create OAuth credentials (Desktop app) → Download JSON → rename to credentials.json")
        print(f"  4. Place credentials.json in: {BASE_DIR}")
        sys.exit(1)

    creds = None
    if TOKEN_FILE.exists():
        creds = Credentials.from_authorized_user_file(str(TOKEN_FILE), SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(str(CREDS_FILE), SCOPES)
            creds = flow.run_local_server(port=0)
        TOKEN_FILE.write_text(creds.to_json())

    return build("drive", "v3", credentials=creds)


def get_or_create_folder(service, folder_name):
    """Get or create a Google Drive folder."""
    results = service.files().list(
        q=f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false",
        fields="files(id, name)"
    ).execute()
    files = results.get("files", [])
    if files:
        return files[0]["id"]
    # Create folder
    meta = {"name": folder_name, "mimeType": "application/vnd.google-apps.folder"}
    folder = service.files().create(body=meta, fields="id").execute()
    return folder["id"]


def upload_docx_as_gdoc(service, docx_path: Path, folder_id: str, doc_name: str):
    """Upload a DOCX file and convert it to Google Docs format."""
    media = MediaFileUpload(
        str(docx_path),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=True
    )
    meta = {
        "name": doc_name,
        "mimeType": "application/vnd.google-apps.document",
        "parents": [folder_id]
    }
    gdoc = service.files().create(body=meta, media_body=media, fields="id, webViewLink").execute()
    return gdoc["id"], gdoc["webViewLink"]


def make_public(service, file_id):
    """Make a Google Doc publicly viewable (anyone with the link)."""
    permission = {"type": "anyone", "role": "reader"}
    service.permissions().create(fileId=file_id, body=permission).execute()


# ─── Main ─────────────────────────────────────────────────────────────────────
def main():
    print("═" * 60)
    print("SOW Google Docs Creator")
    print("═" * 60)

    # Find all fixed markdown files
    md_files = sorted(FIXED_DIR.glob("*.md"))
    if not md_files:
        print(f"❌ No .md files found in {FIXED_DIR}")
        sys.exit(1)

    print(f"Found {len(md_files)} fixed markdown files.\n")

    # Load existing links
    links = {}
    if LINKS_FILE.exists():
        links = json.loads(LINKS_FILE.read_text(encoding="utf-8"))

    # Convert all to DOCX first
    print("Step 1: Converting markdown files to DOCX...")
    docx_files = {}
    for md_file in md_files:
        docx_path = DOCX_DIR / (md_file.stem + ".docx")
        if docx_path.exists():
            print(f"  ✓ {md_file.name} (already exists, skipping conversion)")
        else:
            print(f"  → Converting {md_file.name}...")
            try:
                md_to_docx(md_file, docx_path)
                print(f"  ✓ {docx_path.name}")
            except Exception as e:
                print(f"  ❌ Error: {e}")
                continue
        docx_files[md_file.name] = docx_path

    print(f"\nStep 2: Uploading to Google Drive...")
    print("(A browser window will open for Google authentication if needed)\n")

    try:
        service = get_drive_service()
    except SystemExit:
        return

    folder_id = get_or_create_folder(service, DRIVE_FOLDER_NAME)
    print(f"Drive folder: '{DRIVE_FOLDER_NAME}' (id: {folder_id})\n")

    newly_created = 0
    for md_name, docx_path in docx_files.items():
        # Use filename (without .md) as the document name
        doc_name = md_name.replace(".md", "")
        if doc_name in links or md_name in links:
            print(f"  ✓ {doc_name} — already in gdoc_links.json, skipping")
            continue
        try:
            print(f"  ↑ Uploading: {doc_name}")
            file_id, view_link = upload_docx_as_gdoc(service, docx_path, folder_id, doc_name)
            make_public(service, file_id)
            links[doc_name] = view_link
            # Also store by md_name key
            links[md_name] = view_link
            print(f"    ✅ {view_link}")
            newly_created += 1
        except Exception as e:
            print(f"    ❌ Error uploading {doc_name}: {e}")

    # Save updated links
    LINKS_FILE.write_text(json.dumps(links, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"\n{'═' * 60}")
    print(f"Done! {newly_created} new Google Docs created.")
    print(f"Links saved to: {LINKS_FILE}")
    print(f"\nRefresh the website to see Google Doc links appear on each document card.")
    print(f"{'═' * 60}")


if __name__ == "__main__":
    main()
